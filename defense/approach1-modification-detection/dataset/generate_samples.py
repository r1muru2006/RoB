#!/usr/bin/env python3
"""
Generate benign and malicious file samples for Approach 1 classifier training.

Based on Appendix C of the USENIX Security 2023 RoB paper:
- 5 file types: txt, pdf, docx, xlsx, jpeg
- Benign modifications: content removal + content addition
- Malicious modifications: AES-256-GCM full encryption

Output structure:
  samples/
    originals/{txt,pdf,docx,xlsx,jpeg}/
    benign/{txt,pdf,docx,xlsx,jpeg}/
    malicious/{txt,pdf,docx,xlsx,jpeg}/
"""

import os
import sys
import random
import string
import struct
import io
import shutil
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from shared.crypto_utils import aes_encrypt

SAMPLES_DIR = Path(__file__).resolve().parent / "samples"
FILE_TYPES = ["txt", "pdf", "docx", "xlsx", "jpeg"]
NUM_ORIGINALS = 50
BENIGN_MODS_PER_FILE = 50
WORD_POOL = None


def get_word_pool():
    global WORD_POOL
    if WORD_POOL is None:
        WORD_POOL = [
            ''.join(random.choices(string.ascii_lowercase, k=random.randint(3, 10)))
            for _ in range(5000)
        ]
    return WORD_POOL


def random_text(min_words=200, max_words=1000):
    pool = get_word_pool()
    n = random.randint(min_words, max_words)
    words = [random.choice(pool) for _ in range(n)]
    lines = []
    i = 0
    while i < len(words):
        line_len = random.randint(5, 15)
        lines.append(' '.join(words[i:i+line_len]))
        i += line_len
    return '\n'.join(lines)


# ─── TXT ─────────────────────────────────────────────────────────────────

def generate_txt_original():
    return random_text(300, 800).encode('utf-8')


def modify_txt_benign_remove(data: bytes) -> bytes:
    text = data.decode('utf-8')
    words = text.split()
    if len(words) < 2:
        return data
    r = random.randint(1, max(1, len(words) // 2))
    return ' '.join(words[:-r]).encode('utf-8')


def modify_txt_benign_add(data: bytes) -> bytes:
    pool = get_word_pool()
    r = random.randint(10, 100)
    addition = ' '.join(random.choice(pool) for _ in range(r))
    return data + ('\n' + addition).encode('utf-8')


# ─── PDF ─────────────────────────────────────────────────────────────────

def generate_pdf_original():
    from fpdf import FPDF
    pdf = FPDF()
    num_pages = random.randint(2, 5)
    for _ in range(num_pages):
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        text = random_text(50, 200)
        pdf.multi_cell(0, 5, text)
    return pdf.output()


def modify_pdf_benign_remove(data: bytes) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    text = random_text(30, 100)
    pdf.multi_cell(0, 5, text)
    return pdf.output()


def modify_pdf_benign_add(data: bytes) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    num_pages = random.randint(3, 7)
    for _ in range(num_pages):
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        text = random_text(80, 300)
        pdf.multi_cell(0, 5, text)
    return pdf.output()


# ─── DOCX ────────────────────────────────────────────────────────────────

def generate_docx_original():
    from docx import Document
    doc = Document()
    num_paragraphs = random.randint(5, 20)
    for _ in range(num_paragraphs):
        doc.add_paragraph(random_text(20, 80))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def modify_docx_benign_remove(data: bytes) -> bytes:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = doc.paragraphs
    if len(paragraphs) > 1:
        r = random.randint(1, max(1, len(paragraphs) // 2))
        for _ in range(r):
            if doc.paragraphs:
                p = doc.paragraphs[-1]
                p._element.getparent().remove(p._element)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def modify_docx_benign_add(data: bytes) -> bytes:
    from docx import Document
    doc = Document(io.BytesIO(data))
    r = random.randint(2, 10)
    for _ in range(r):
        doc.add_paragraph(random_text(20, 60))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── XLSX ────────────────────────────────────────────────────────────────

def generate_xlsx_original():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    num_rows = random.randint(50, 200)
    num_cols = random.randint(3, 8)
    for r in range(1, num_rows + 1):
        for c in range(1, num_cols + 1):
            if random.random() < 0.5:
                ws.cell(row=r, column=c, value=random.uniform(-1000, 1000))
            else:
                pool = get_word_pool()
                ws.cell(row=r, column=c, value=random.choice(pool))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def modify_xlsx_benign_remove(data: bytes) -> bytes:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data))
    ws = wb.active
    max_row = ws.max_row
    if max_row > 2:
        r = random.randint(1, max(1, max_row // 2))
        ws.delete_rows(max_row - r + 1, r)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def modify_xlsx_benign_add(data: bytes) -> bytes:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data))
    ws = wb.active
    max_row = ws.max_row
    num_cols = ws.max_column
    r = random.randint(5, 30)
    for row in range(max_row + 1, max_row + r + 1):
        for c in range(1, num_cols + 1):
            if random.random() < 0.5:
                ws.cell(row=row, column=c, value=random.uniform(-1000, 1000))
            else:
                pool = get_word_pool()
                ws.cell(row=row, column=c, value=random.choice(pool))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── JPEG ────────────────────────────────────────────────────────────────

def generate_jpeg_original():
    from PIL import Image
    w = random.randint(200, 800)
    h = random.randint(200, 800)
    img = Image.new('RGB', (w, h))
    pixels = img.load()
    for x in range(w):
        for y in range(h):
            pixels[x, y] = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=85)
    return buf.getvalue()


def modify_jpeg_benign_remove(data: bytes) -> bytes:
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    w, h = img.size
    r1 = random.randint(1, max(1, w - 1))
    r2 = random.randint(1, max(1, h - 1))
    img = img.crop((0, 0, r1, r2))
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=85)
    return buf.getvalue()


def modify_jpeg_benign_add(data: bytes) -> bytes:
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    w, h = img.size
    extra_w = random.randint(10, 100)
    extra_h = random.randint(10, 100)
    new_img = Image.new('RGB', (w + extra_w, h + extra_h))
    new_img.paste(img, (0, 0))
    pixels = new_img.load()
    for x in range(w, w + extra_w):
        for y in range(h + extra_h):
            pixels[x, y] = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
    for y in range(h, h + extra_h):
        for x in range(w):
            pixels[x, y] = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
    buf = io.BytesIO()
    new_img.save(buf, format='JPEG', quality=85)
    return buf.getvalue()


# ─── Generators registry ─────────────────────────────────────────────────

GENERATORS = {
    "txt":  (generate_txt_original, modify_txt_benign_remove, modify_txt_benign_add, ".txt"),
    "pdf":  (generate_pdf_original, modify_pdf_benign_remove, modify_pdf_benign_add, ".pdf"),
    "docx": (generate_docx_original, modify_docx_benign_remove, modify_docx_benign_add, ".docx"),
    "xlsx": (generate_xlsx_original, modify_xlsx_benign_remove, modify_xlsx_benign_add, ".xlsx"),
    "jpeg": (generate_jpeg_original, modify_jpeg_benign_remove, modify_jpeg_benign_add, ".jpeg"),
}


def main():
    random.seed(42)

    if SAMPLES_DIR.exists():
        shutil.rmtree(SAMPLES_DIR)

    for ft in FILE_TYPES:
        (SAMPLES_DIR / "originals" / ft).mkdir(parents=True, exist_ok=True)
        (SAMPLES_DIR / "benign" / ft).mkdir(parents=True, exist_ok=True)
        (SAMPLES_DIR / "malicious" / ft).mkdir(parents=True, exist_ok=True)

    for ft in FILE_TYPES:
        gen_orig, mod_remove, mod_add, ext = GENERATORS[ft]
        print(f"[*] Generating {ft} samples...")

        for i in range(NUM_ORIGINALS):
            orig_data = gen_orig()
            orig_path = SAMPLES_DIR / "originals" / ft / f"{ft}_{i:04d}{ext}"
            orig_path.write_bytes(orig_data)

            half = BENIGN_MODS_PER_FILE // 2
            for j in range(half):
                try:
                    mod_data = mod_remove(orig_data)
                    mod_path = SAMPLES_DIR / "benign" / ft / f"{ft}_{i:04d}_rm_{j:03d}{ext}"
                    mod_path.write_bytes(mod_data)
                except Exception:
                    pass

            for j in range(half):
                try:
                    mod_data = mod_add(orig_data)
                    mod_path = SAMPLES_DIR / "benign" / ft / f"{ft}_{i:04d}_add_{j:03d}{ext}"
                    mod_path.write_bytes(mod_data)
                except Exception:
                    pass

            enc_data = aes_encrypt(orig_data)
            enc_path = SAMPLES_DIR / "malicious" / ft / f"{ft}_{i:04d}_enc{ext}"
            enc_path.write_bytes(enc_data)

            if (i + 1) % 10 == 0:
                print(f"    {ft}: {i+1}/{NUM_ORIGINALS} originals done")

    total_orig = sum(len(list((SAMPLES_DIR / "originals" / ft).iterdir())) for ft in FILE_TYPES)
    total_benign = sum(len(list((SAMPLES_DIR / "benign" / ft).iterdir())) for ft in FILE_TYPES)
    total_mal = sum(len(list((SAMPLES_DIR / "malicious" / ft).iterdir())) for ft in FILE_TYPES)
    print(f"\n[+] Done! Originals: {total_orig}, Benign mods: {total_benign}, Malicious: {total_mal}")


if __name__ == "__main__":
    main()
